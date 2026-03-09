from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.shared import OxmlElement, qn

out_path = "Andrew_Black_Resume.docx"

doc = Document()

# Page setup
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

# Helpers
def set_paragraph_spacing(p, before=0, after=0, line=1.0):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line

def add_section_heading(text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11.5)
    set_paragraph_spacing(p, before=10, after=4, line=1.0)
    return p

def add_bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    set_paragraph_spacing(p, before=0, after=0, line=1.05)
    for r in p.runs:
        r.font.size = Pt(10.5)
    return p

def add_job_header(left_text, right_text):
    p = doc.add_paragraph()
    # Right-aligned tab stop at right margin
    tabstops = p.paragraph_format.tab_stops
    tabstops.clear_all()
    tabstops.add_tab_stop(Inches(6.9), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
    run_left = p.add_run(left_text + "\t")
    run_left.bold = True
    run_left.font.size = Pt(11)
    run_right = p.add_run(right_text)
    run_right.font.size = Pt(10.5)
    set_paragraph_spacing(p, before=2, after=1, line=1.0)
    return p

def add_subline(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10.5)
    set_paragraph_spacing(p, before=0, after=2, line=1.0)
    return p

# Header
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Andrew Black")
run.bold = True
run.font.size = Pt(20)
set_paragraph_spacing(p, before=0, after=2, line=1.0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact = "NJ [USA] | (215) 375-8715 | ab4@comcast.net | linkedin.com/in/andrewablack | github.com/ablack94"
run = p.add_run(contact)
run.font.size = Pt(10.5)
set_paragraph_spacing(p, before=0, after=8, line=1.0)

# Summary
add_section_heading("Summary")
summary = (
    "Senior Software Engineer with deep systems fundamentals and a growing focus on agentic AI infrastructure. "
    "At Bloomberg I serve as an embedded technical lead across multiple engineering teams, designing cross-cutting platform work and building "
    "agentic tooling with emphasis on sandboxing, auditability, and safe autonomous execution. "
    "Independently building containerized multi-agent environments with network isolation, structured evaluation, and audit trails."
)
p = doc.add_paragraph(summary)
set_paragraph_spacing(p, before=0, after=0, line=1.15)
for r in p.runs:
    r.font.size = Pt(10.75)

# Core Tech
p = doc.add_paragraph()
set_paragraph_spacing(p, before=6, after=0, line=1.05)
run = p.add_run("Core tech: ")
run.bold = True
run.font.size = Pt(10.5)
run2 = p.add_run("Rust, Python, C++, Linux, Bash | Distributed systems, containerization, sandboxing, networking, concurrency, performance profiling | Agentic AI, MCP, LLM tooling, evaluation design")
run2.font.size = Pt(10.5)

# Experience
add_section_heading("Experience")

# Bloomberg LP
add_job_header("Bloomberg LP | Sr. Software Engineer (Core Infrastructure CHAMP)", "Aug 2018 – Present")
add_subline("Princeton, NJ")
add_bullet("Embedded across multiple engineering teams as a Core Infrastructure CHAMP, providing sustained technical leadership from design through production on platform reliability, API contracts, and cross-service integration.")
add_bullet("Lead cross-stack architecture initiatives spanning auth, observability, rate limiting, and service design; define interfaces and operational standards adopted across the organization.")
add_bullet("Design and build agentic AI workflows using Claude and Model Context Protocol (MCP) servers, with emphasis on sandboxed execution, safety guardrails, auditability, and predictable agent behavior.")
add_bullet("Champion the adoption of MCP beyond internal developer tooling into product-facing integrations, partnering with teams to design interfaces and permission models that bring AI capabilities to engineering and data partners.")
add_bullet("Scale knowledge through design reviews, technical talks, and hands-on workshops; build lasting cross-team relationships that outlast individual project engagements.")
add_bullet("Serve as a go-to diagnostician for complex production incidents and performance regressions; routinely pulled into unfamiliar codebases to identify root causes and deliver pragmatic fixes.")

# Defense contracting
add_job_header("Gnostech Inc. (Lockheed Martin subcontract) | Software Engineer", "Feb 2016 – Aug 2018")
add_subline("Moorestown, NJ")
add_bullet("Cut report generation time from days/weeks to hours by designing and implementing an automated data collection, analysis, and reporting framework (Python/Pandas/NumPy/Matplotlib + C++/Java + Django/Bash).")
add_bullet("Troubleshot application and Linux kernel-level performance issues using tools such as perf, strace, gprof, and procfs to ensure performance deliverables were met.")
add_bullet("Prepared and presented live customer demos; incorporated feedback quickly to iterate on features and usability in a fast-paced Agile environment.")

# SiriusXM (2nd Drexel co-op)
add_job_header("SiriusXM | Software Engineer (Co-op)", "Mar 2014 – Sep 2014")
add_subline("Lawrenceville, NJ")
add_bullet("Designed and implemented real-time monitoring systems for multimedia streams to ensure high-quality satellite radio broadcasts.")
add_bullet("Built a high-uptime redundant pipeline to process multimedia content through a third-party speech-to-text service (Ruby, AWS, PostgreSQL).")

# Anamir Electronics (1st Drexel co-op)
add_job_header("Anamir Electronics | Electrical Engineer (Co-op)", "Sep 2012 – Mar 2013")
add_subline("Yardley, PA")
add_bullet("Wrote software (Python/APL) to procedurally generate PCB trace patterns in the GERBER file format, enabling rapid prototyping of new printed circuit board designs.")
add_bullet("Maintained and repaired precision lab equipment including oscilloscopes, multimeters, and function generators.")

# Education
add_section_heading("Education")
add_job_header("Drexel University", "Philadelphia, PA")
add_subline("B.S. & M.S., Computer Science — Concentrations: Artificial Intelligence, Operating Systems — Jun 2015")
p = doc.add_paragraph()
set_paragraph_spacing(p, before=0, after=1, line=1.05)
run = p.add_run("GPA 3.64 | Major GPA 3.77 | Upsilon Pi Epsilon | Dean's List (2013–2015)")
run.font.size = Pt(10.5)

doc.save(out_path)

out_path
